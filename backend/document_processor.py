import os
import io
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .vision_ocr import IMAGE_EXTENSIONS, extract_text_from_image, extract_text_from_pdf_page_image

# Minimum characters per PDF page to consider it "text-bearing"
# Pages below this threshold are treated as scanned/image-only
MIN_TEXT_PER_PAGE = 30


def extract_text_from_bytes(file_bytes: bytes, filename: str, api_key: str = "") -> dict:
    """
    Extracts text from uploaded file bytes based on the filename extension.
    Supports txt, pdf, docx, and image files (jpg, png, webp, heic, etc.).
    Handles encoding fallbacks, table extractions, and scanned PDF detection.

    Returns a dict with:
        - text: the extracted text content
        - ocr_used: True if vision OCR was used for extraction
    """
    _, extension = os.path.splitext(filename)
    extension = extension.lower()

    # --- Image files: use Groq Vision OCR ---
    if extension in IMAGE_EXTENSIONS:
        if not api_key:
            raise ValueError(
                "Image text extraction requires a Groq API key. "
                "Please configure your key in the Server Status settings."
            )
        text = extract_text_from_image(file_bytes, filename, api_key)
        return {"text": text, "ocr_used": True}

    # --- Plain text files ---
    if extension == '.txt':
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1 encoding if UTF-8 fails
                text = file_bytes.decode("latin-1")
            except Exception as e:
                raise ValueError(f"Failed to decode text file: {str(e)}")
        return {"text": text, "ocr_used": False}

    # --- PDF files (with scanned page fallback) ---
    elif extension == '.pdf':
        try:
            pdf_stream = io.BytesIO(file_bytes)
            pdf_reader = PdfReader(pdf_stream)
            text_parts = []
            ocr_used = False
            ocr_page_count = 0

            for page_idx, page in enumerate(pdf_reader.pages):
                # Try standard text extraction first
                extracted = page.extract_text()

                if extracted and len(extracted.strip()) >= MIN_TEXT_PER_PAGE:
                    # Page has sufficient text — use it directly
                    text_parts.append(extracted)
                elif api_key:
                    # Page appears to be scanned/image-only — try vision OCR
                    print(f">>> [DOC] Page {page_idx + 1}: only {len((extracted or '').strip())} chars, "
                          f"attempting vision OCR fallback...")
                    try:
                        ocr_text = extract_text_from_pdf_page_image(page, page_idx, api_key)
                        if ocr_text and ocr_text.strip():
                            text_parts.append(ocr_text)
                            ocr_used = True
                            ocr_page_count += 1
                        elif extracted:
                            # Vision OCR returned nothing, keep whatever text we had
                            text_parts.append(extracted)
                    except ValueError as ve:
                        # API/auth error — add whatever text we had and log
                        print(f">>> [DOC] Page {page_idx + 1} vision OCR failed: {ve}")
                        if extracted:
                            text_parts.append(extracted)
                elif extracted:
                    # No API key available — keep whatever partial text we got
                    text_parts.append(extracted)

            full_text = "\n".join(text_parts)

            if ocr_used:
                print(f">>> [DOC] PDF processed: {len(pdf_reader.pages)} pages total, "
                      f"{ocr_page_count} pages via vision OCR")

            return {"text": full_text, "ocr_used": ocr_used}

        except PdfReadError as pre:
            raise ValueError(f"Corrupted or encrypted PDF document: {str(pre)}")
        except ValueError:
            raise  # Re-raise vision OCR errors
        except Exception as e:
            raise ValueError(f"Failed to parse PDF document: {str(e)}")

    # --- Word documents ---
    elif extension == '.docx':
        try:
            docx_stream = io.BytesIO(file_bytes)
            doc = Document(docx_stream)

            # Extract paragraphs
            paragraphs_text = [para.text for para in doc.paragraphs]

            # Extract tables (ensures text inside Word tables isn't ignored)
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    filtered_row = [t.strip() for t in row_text if t and t.strip()]
                    if filtered_row:
                        tables_text.append(" | ".join(filtered_row))

            combined_text = "\n".join(paragraphs_text + tables_text)
            return {"text": combined_text, "ocr_used": False}
        except Exception as e:
            raise ValueError(f"Failed to parse Word (.docx) document: {str(e)}")

    else:
        raise ValueError(f"Unsupported file extension: {extension}")

def get_text_chunks(text: str, chunk_size: int = 1000, chunk_overlap: int = 200):
    """
    Splits a large text string into smaller chunks using LangChain's RecursiveCharacterTextSplitter.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    chunks = text_splitter.split_text(text)
    return chunks
